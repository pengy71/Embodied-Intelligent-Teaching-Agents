from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(__file__).parent

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = bold; r.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True); shade(t.rows[0].cells[i], 'F2F4F7')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row): set_cell_text(cells[i], str(val))
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def base(title, subtitle):
    d = Document(); s = d.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.8)
    normal = d.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
    for name, size, color in [('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5'),('Heading 3',12,'1F4D78')]:
        st=d.styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(title); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(11,37,69)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(subtitle); r.font.size=Pt(11); r.font.color.rgb=RGBColor(89,89,89)
    d.add_paragraph('版本：v0.6  |  日期：2026-07-31  |  项目：基于 OpenMAIC 的具身智能课程教学智能体')
    return d

def architecture():
    d=base('技术架构文档 v0.6','系统模块、技术选型、核心接口与数据流说明')
    d.add_heading('1. 架构目标',1); d.add_paragraph('在 OpenMAIC 多智能体课堂框架基础上，增加课程知识层、学生学习层和教师教学支持层。本版本重点落地学生端答疑中心、练习测试，以及教师端数据反馈所需的基础数据闭环。')
    d.add_heading('2. 总体架构',1)
    table(d,['层级','主要模块','本版本职责'],[
        ('交互层','教师端、学生端、课堂端','答疑中心、练习中心、课堂互动与报告入口'),
        ('业务层','课程知识、答疑、练习、学习记录','问题路由、题目生成、评分、错题与掌握度'),
        ('智能体层','课程知识智能体、学习规划智能体、教师助手','知识检索、回答组织、练习推荐与教学建议'),
        ('数据层','知识库、知识图谱、RuntimeStore/数据库','课程资料、知识点、问答记录、答题记录、错题集'),
        ('基础设施','Next.js、LangGraph、LLM、对象存储','页面渲染、智能体编排、模型调用、文件管理'),
    ],[1.1,2.0,3.3])
    d.add_heading('3. 核心技术选型',1)
    table(d,['技术','选型','原因'],[
        ('前端','Next.js + React + Tailwind','复用队友 OpenMAIC 页面和组件体系'),
        ('智能体编排','LangGraph','复用已有多智能体导演图与对话循环'),
        ('课程问答','RAG + 知识图谱检索','支持课程限定回答、关联知识点和教材溯源'),
        ('练习评分','本地客观题评分 + LLM 简答题评分','低延迟且兼顾开放题评价'),
        ('持久化','RuntimeStore；后续接 PostgreSQL/D1','保存答题尝试、错题和学习轨迹'),
    ],[1.2,2.0,3.2])
    d.add_heading('4. 答疑中心数据流',1)
    d.add_paragraph('学生输入问题 → 问题分类 → 课程知识检索 → 课程知识智能体组织答案 → 返回教材页码/段落 → 关联知识点 → 推荐专项练习 → 记录问答轨迹。')
    d.add_heading('5. 练习测试数据流',1)
    d.add_paragraph('选择练习类型 → 根据章节、知识点、难度生成题目 → 学生作答 → 客观题本地判分/简答题调用 /api/quiz-grade → 汇总成绩 → 更新知识点掌握度 → 错题入库 → 推荐下一轮练习。')
    d.add_heading('6. 核心接口',1)
    table(d,['接口','方法','输入','输出'],[
        ('/api/course/ask','POST','courseId、question、knowledgePointId','answer、citations、relatedKnowledge、recommendedExercises'),
        ('/api/exercises/generate','POST','mode、chapterId、knowledgePointIds、difficulty、count','QuizQuestion[]'),
        ('/api/quiz-grade','POST','question、userAnswer、points','score、comment'),
        ('/api/exercises/submit','POST','attemptId、answers','score、knowledgeMastery、wrongAnswers'),
        ('/api/exercises/wrong-answers','GET','studentId、courseId','错题列表与复习状态'),
    ],[1.7,0.7,2.4,2.0])
    d.add_heading('7. 安全与边界',1); d.add_paragraph('课程问答必须优先使用课程资料；无法检索到依据时明确提示“课程资料未覆盖”，不得伪造页码。学生数据按 studentId 隔离；模型密钥只在服务端使用；练习答案不在客户端接口中提前暴露。')
    d.add_heading('8. 本版本验收标准',1)
    for x in ['能提出课程问题并返回结构化回答、引用和相关知识点。','能从答疑结果一键生成 3—5 道专项练习。','客观题可自动判分，简答题可调用 AI 评分。','提交后显示成绩、解析、薄弱知识点和错题记录。','现有 OpenMAIC 课堂生成与播放流程不受影响。']:
        d.add_paragraph(x, style='List Bullet')
    d.save(OUT/'技术架构文档v0.6-具身智能课程教学智能体.docx')

def prd():
    d=base('产品需求文档 PRD v0.6','答疑中心与练习测试模块')
    d.add_heading('1. 需求背景',1); d.add_paragraph('具身智能课程知识跨度大、资料分散，学生在学习过程中需要可靠答疑、教材出处、前置知识和针对性练习。本版本以“提问—理解—练习—反馈—再学习”为核心闭环。')
    d.add_heading('2. 产品目标',1)
    table(d,['目标','衡量方式'],[('降低学生查找课程资料的成本','问题可在一个入口完成，回答带课程来源'),('提高练习针对性','练习绑定章节/知识点/难度'),('形成学习数据闭环','提交后生成掌握度和错题记录')],[3.2,3.2])
    d.add_heading('3. 功能范围',1)
    table(d,['功能模块','功能点','优先级'],[
        ('答疑中心','快速提问、历史问答、教材溯源、相关知识点、生成专项练习','P0'),
        ('练习测试','AI生成练习、章节练习、专项练习、阶段测试','P0'),
        ('结果反馈','自动判分、答案解析、薄弱知识点、学习建议','P0'),
        ('错题集','错题保存、按知识点筛选、重新练习、标记掌握','P1'),
        ('个性化设置','教学风格、内容深度、练习难度偏好','P1'),
    ],[1.3,4.5,0.7])
    d.add_heading('4. 关键交互流程',1)
    d.add_heading('4.1 快速提问',2)
    for x in ['学生进入“答疑中心”，输入问题并点击“提问”。','系统检索课程知识库和知识图谱。','AI返回解释、教材页码/段落、相关知识点。','学生可点击“生成相关专项练习”，进入练习页面。']:
        d.add_paragraph(x, style='List Number')
    d.add_heading('4.2 练习与反馈',2)
    for x in ['学生选择 AI生成、章节、专项、测试或错题集。','系统生成题目，显示题型、难度、知识点和预计用时。','学生完成所有题目后提交。','客观题自动判分，简答题由 AI 评分。','系统展示分数、解析、薄弱知识点，并将错误题目写入错题集。']:
        d.add_paragraph(x, style='List Number')
    d.add_heading('5. 页面说明',1)
    table(d,['页面','核心组件','状态'],[('答疑中心','输入框、AI回答卡片、引用卡片、知识点标签、推荐练习按钮','空状态/回答中/回答完成/失败重试'),('练习中心','练习类型、题目卡片、选项、提交按钮','未开始/答题中/提交中/结果页'),('错题集','筛选器、错题卡片、重新练习按钮','无错题/有错题/已掌握')],[1.4,3.8,1.4])
    d.add_heading('6. 示例内容',1); d.add_paragraph('示例问题：视觉感知和世界模型有什么区别？\n预期回答：视觉感知负责获取环境信息，世界模型负责对环境状态进行表示、推理和预测。\n预期操作：显示教材第86页，关联“环境感知、世界模型、任务规划”，点击后生成5道专项练习。')
    d.add_heading('7. 非功能需求',1)
    for x in ['回答首屏反馈不超过 5 秒；生成练习过程显示加载状态。','回答引用必须来自课程资料，无法溯源时明确提示。','移动端和桌面端均可正常使用；关键按钮保持明显。','所有答题记录支持刷新后恢复，重复提交不产生重复记录。']:
        d.add_paragraph(x, style='List Bullet')
    d.add_heading('8. 版本计划',1)
    table(d,['版本','范围'],[('v0.6（本次）','答疑中心、专项练习、自动判分、结果解析、基础错题记录'),('v0.7','接入真实课程知识库、教材页码溯源、章节练习和历史问答'),('v0.8','学习画像、个性化推荐、教师端班级学情分析')],[1.5,5.1])
    d.save(OUT/'PRDv0.6-答疑中心与练习测试.docx')

if __name__=='__main__':
    architecture(); prd()
